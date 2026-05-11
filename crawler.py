from __future__ import annotations

import hashlib
import heapq
import json
import math
import re
import time
from collections import Counter
from collections.abc import Callable, Iterable
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import suppress
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, cast
from urllib.parse import parse_qsl

import requests
from bs4 import BeautifulSoup, Tag
from docx import Document as create_doc
from docx.document import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.worksheet import Worksheet

from export.csv_exporter import build_export_row as _export_build_row
from export.csv_exporter import export_posts_csv as _export_posts_csv
from export.markdown_exporter import export_weekly_report_md as _export_weekly_report_md
from export.summary_exporter import build_summary as _export_build_summary
from export.summary_exporter import calc_date_distribution_fit as _export_calc_date_distribution_fit
from export.summary_exporter import write_summary_txt as _export_write_summary_txt
from modules.crawler_filters import should_exclude_post
from modules.crawler_scoring import calculate_score
from modules.text_cleaning import clean_topic_tags, collapse_blank_lines, normalize_weibo_text, strip_html_text
from modules.time_utils import normalize_date as _normalize_date
from modules.time_utils import parse_weibo_time
from modules.weibo_url import normalize_image_url, parse_super_topic_id as _parse_super_topic_id, to_absolute_url

FM_VIEW_MARKER = "FM.view("
COMMENTS_API_URL = "https://weibo.com/ajax/statuses/buildComments"
COMMENT_ANALYSIS_MIN_ROWS = 60
COMMENT_ANALYSIS_RATIO = 0.36
FINAL_COMMENT_ANALYSIS_LIMIT = 45
FORWARDED_NODE_TYPES = {
    "feed_list_forwardContent",
    "feed_list_forwardContent_full",
}
FORWARDED_CLASS_NAMES = {
    "WB_feed_expand",
    "WB_feed_expand_v2",
    "WB_expand",
    "WB_feed_expand_media",
}
EXPORT_COLUMN_MAP = [
    ("user_name", "作者昵称"),
    ("publish_time", "帖子发送时间"),
    ("post_url", "帖子链接"),
    ("original_image_urls", "原图链接"),
    ("image_local_paths", "图片本地路径"),
    ("content", "帖子内容"),
    ("reposts", "转发数"),
    ("comments", "评论总数"),
    ("likes", "点赞数"),
    ("non_author_comments", "非楼主评论数"),
    ("author_replies", "楼主回复数"),
    ("topic_comment_factor", "话题评论系数"),
    ("score", "帖子分数"),
    ("engagement_total", "互动总量"),
    ("top_comment_1", "热评1(按点赞)"),
    ("top_comment_2", "热评2(按点赞)"),
    ("top_comment_3", "热评3(按点赞)"),
]
EMBED_IMAGE_HEADER_PREFIX = "图片预览"
WIDE_COLUMNS = {
    "帖子发送时间": 20,
    "帖子链接": 45,
    "原图链接": 80,
    "图片本地路径": 80,
    "帖子内容": 60,
    "热评1(按点赞)": 65,
    "热评2(按点赞)": 65,
    "热评3(按点赞)": 65,
}
DOCX_SIZE_LIMIT_BYTES = 10 * 1000 * 1000


class CrawlError(Exception):
    pass


@dataclass
class CrawlConfig:
    super_topic: str
    cookie: str
    max_pages: int = 30
    pause_seconds: float = 1.0
    days_window: int = 7
    topic_comment_factor: float = 1.0
    comment_page_limit: int = 8
    text_workers: int = 6
    comment_workers: int = 6
    window_start: datetime | None = None
    window_end: datetime | None = None
    carryover_hours: int = 0


class WeiboSuperTopicCrawler:
    def __init__(
        self,
        cookie: str,
        user_agent: str | None = None,
        progress_callback: Callable[[str], None] | None = None,
        stage_callback: Callable[[str, list[dict]], None] | None = None,
        comment_cache_reader: Callable[[str], dict | None] | None = None,
        comment_cache_writer: Callable[[str, dict], None] | None = None,
    ) -> None:
        self.cookie = cookie.strip()
        self.user_agent = user_agent or (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        )
        self.progress_callback = progress_callback
        self.stage_callback = stage_callback
        self.comment_cache_reader = comment_cache_reader
        self.comment_cache_writer = comment_cache_writer
        self.topic_name = ""
        self.report_title = "微博超话周报"
        self.session = requests.Session()
        self.session.headers.update(
            {
                "User-Agent": self.user_agent,
                "Cookie": self.cookie,
                "X-Requested-With": "XMLHttpRequest",
                "Accept": "application/json, text/plain, */*",
            }
        )

    def _new_session(self) -> requests.Session:
        session = requests.Session()
        session.headers.update(self.session.headers)
        return session

    def _log(self, message: str) -> None:
        if self.progress_callback:
            self.progress_callback(message)

    def _emit_stage_cache(self, stage: str, posts: list[dict]) -> None:
        if self.stage_callback:
            self.stage_callback(stage, posts)

    def crawl(self, config: CrawlConfig) -> list[dict]:
        topic_id = parse_super_topic_id(config.super_topic)
        if not topic_id:
            raise CrawlError("无法从输入中解析超话ID，请确认链接格式。")
        if not self.cookie:
            raise CrawlError("Cookie 为空。请先从已登录的微博页面复制 Cookie。")

        referer = f"https://weibo.com/p/{topic_id}/super_index"
        use_fixed_window = bool(config.window_start and config.window_end)
        if use_fixed_window:
            window_start = config.window_start or (datetime.now() - timedelta(days=config.days_window))
            window_end = config.window_end or datetime.now()
            if window_end <= window_start:
                raise CrawlError("固定周窗口参数错误：结束时间必须晚于开始时间。")
            self._log(
                f"按固定窗口抓取：{window_start.strftime('%Y-%m-%d %H:%M')} 至 "
                f"{window_end.strftime('%Y-%m-%d %H:%M')}"
            )
        else:
            window_end = datetime.now()
            window_start = window_end - timedelta(days=config.days_window)
            self._log(f"按近 {config.days_window} 天抓取：{window_start.strftime('%Y-%m-%d %H:%M')} 起")

        carryover_hours = max(0, int(config.carryover_hours or 0))
        if carryover_hours > 0:
            self._log(
                f"已启用统计保险窗口：截止前 {carryover_hours} 小时发布的帖子顺延到下一期。"
            )

        all_posts: list[dict] = []
        seen_post_ids: set[str] = set()
        last_page_sig: tuple[str, ...] | None = None
        consecutive_same_pages = 0
        no_hit_streak = 0

        for page in range(1, config.max_pages + 1):
            self._log(f"抓取第 {page} 页...")
            page_html = self._fetch_super_index_page(referer, page)
            if not self.topic_name:
                self.topic_name = extract_super_topic_name(page_html, config.super_topic)
                self.report_title = build_report_title(self.topic_name, config.super_topic)
                if self.topic_name:
                    self._log(f"已识别超话名称：{self.topic_name}")
            feed_html = extract_feed_html_from_page(page_html)
            page_posts = parse_posts_from_html(feed_html)
            if not page_posts:
                self._log("本页没有帖子数据，停止翻页。")
                break

            # 用整页签名做“连续重复页”判断，避免仅比较前几条导致误判提前停页
            sig = tuple(
                str(post.get("post_id", "")).strip()
                for post in page_posts
                if str(post.get("post_id", "")).strip()
            )
            if sig and last_page_sig == sig:
                consecutive_same_pages += 1
                self._log(f"检测到重复页内容（连续{consecutive_same_pages + 1}页相同），继续翻页...")
            else:
                consecutive_same_pages = 0
            last_page_sig = sig if sig else last_page_sig

            page_recent = 0
            for post in page_posts:
                post_id = str(post.get("post_id") or "").strip()
                post_dt = parse_publish_datetime(str(post.get("publish_time") or ""))
                bucket_dt = _apply_carryover_bucket(post_dt, carryover_hours)

                if bucket_dt and window_start <= bucket_dt < window_end:
                    page_recent += 1

                keep = bucket_dt is not None and window_start <= bucket_dt < window_end
                if not keep:
                    continue
                if post_id and post_id in seen_post_ids:
                    continue
                if post_id:
                    seen_post_ids.add(post_id)
                all_posts.append(post)

            if use_fixed_window:
                self._log(f"第 {page} 页读取 {len(page_posts)} 条，窗口内命中 {page_recent} 条。")
            else:
                self._log(f"第 {page} 页读取 {len(page_posts)} 条，近 {config.days_window} 天命中 {page_recent} 条。")

            # 新停页规则：只有连续5页都没有命中时间窗口，才停止翻页
            if page_recent == 0:
                no_hit_streak += 1
                self._log(f"连续无命中页：{no_hit_streak}/5")
            else:
                no_hit_streak = 0

            if no_hit_streak >= 5:
                self._log("已连续5页无时间窗口命中帖子，停止翻页。")
                break

            time.sleep(max(config.pause_seconds, 0))

        if not all_posts:
            if use_fixed_window:
                raise CrawlError(
                    "未抓到该固定窗口内帖子。可能是 Cookie 失效，或该时段无数据。"
                )
            raise CrawlError(
                f"未抓到近 {config.days_window} 天帖子。可能是 Cookie 失效，或超话近 {config.days_window} 天无数据。"
            )

        self._emit_stage_cache("posts_raw", all_posts)
        self._log("补全帖子正文（包含疑似截断内容）...")
        self.hydrate_full_text_posts(all_posts, max_workers=config.text_workers)
        self._emit_stage_cache("posts_hydrated", all_posts)

        self._log("开始计算评分（包含评论结构估算与时间权重）...")
        self.enrich_score_fields(all_posts, config)

        self._log("自动校准时间权重（目标拟合度 90%~93%）...")
        self._recalibrate_time_weight(all_posts, config)
        for _ in range(2):
            if not self._ensure_candidate_comment_analysis(all_posts, config):
                break
            self._recalibrate_time_weight(all_posts, config)

        all_posts.sort(key=lambda x: float(x.get("score", 0.0)), reverse=True)
        self._emit_stage_cache("posts_scored", all_posts)
        return all_posts

    def _fetch_super_index_page(self, super_index_url: str, page: int) -> str:
        resp = self.session.get(
            super_index_url,
            params={"page": str(page)},
            headers={"Referer": "https://weibo.com/"},
            timeout=20,
        )
        text = resp.text.strip()
        if "<title>Sina Visitor System</title>" in text:
            raise CrawlError("微博返回访客验证页面。请使用已登录账号 Cookie 重试。")
        if resp.status_code >= 400:
            raise CrawlError(f"加载超话页面失败，HTTP {resp.status_code}")
        return text

    def enrich_score_fields(self, posts: list[dict], config: CrawlConfig) -> None:
        rows = list(posts)
        total = len(rows)
        if not total:
            return

        for post in rows:
            self._set_estimated_score_fields(post, config)

        comment_rows = [
            post
            for post in rows
            if int(post.get("comments", 0) or 0) > 0
            and str(post.get("post_id") or "").strip()
            and str(post.get("author_id") or "").strip()
        ]
        analysis_rows = self._select_comment_analysis_rows(comment_rows, config)
        analysis_total = len(analysis_rows)
        skipped = len(comment_rows) - analysis_total
        if skipped > 0:
            self._log(f"快速评分已完成：评论结构精查 {analysis_total}/{len(comment_rows)} 条，跳过低潜力 {skipped} 条。")
        elif comment_rows:
            self._log(f"快速评分已完成：评论结构精查 {analysis_total}/{len(comment_rows)} 条。")

        if not analysis_rows:
            return

        worker_count = _bounded_worker_count(config.comment_workers, analysis_total)
        if worker_count == 1:
            for idx, post in enumerate(analysis_rows, start=1):
                self._log(f"评分进度 {idx}/{analysis_total}: {post.get('post_id', '-')}")
                self._enrich_score_fields(post, config)
            return

        self._log(f"评论结构与基础评分并行处理：{worker_count} 个线程。")
        with ThreadPoolExecutor(max_workers=worker_count, thread_name_prefix="score-enrich") as executor:
            futures = {
                executor.submit(self._enrich_score_fields_with_private_session, post, config): post
                for post in analysis_rows
            }
            for completed, future in enumerate(as_completed(futures), start=1):
                post = futures[future]
                try:
                    future.result()
                except Exception as err:
                    self._log(
                        f"评分失败 {completed}/{analysis_total}: {post.get('post_id', '-')}, "
                        f"{type(err).__name__}: {err}"
                    )
                    self._set_estimated_score_fields(post, config)
                if completed == analysis_total or completed % 5 == 0:
                    self._log(f"评分进度 {completed}/{analysis_total}")

    def _select_comment_analysis_rows(self, rows: list[dict], config: CrawlConfig) -> list[dict]:
        if not rows:
            return []
        total = len(rows)
        if total <= COMMENT_ANALYSIS_MIN_ROWS:
            return rows
        limit = min(
            total,
            max(
                COMMENT_ANALYSIS_MIN_ROWS,
                int(total * COMMENT_ANALYSIS_RATIO),
                FINAL_COMMENT_ANALYSIS_LIMIT,
            ),
        )
        ref_now = config.window_end or datetime.now()
        return heapq.nlargest(
            limit,
            rows,
            key=lambda post: self._estimated_priority_score(post, config, ref_now),
        )

    def _estimated_priority_score(self, post: dict, config: CrawlConfig, ref_now: datetime) -> float:
        likes = int(post.get("likes", 0) or 0)
        comments = int(post.get("comments", 0) or 0)
        reposts = int(post.get("reposts", 0) or 0)
        comment_factor = max(0.5, float(config.topic_comment_factor))
        base = likes * 0.3 + comments * 0.5 * comment_factor + reposts * 0.1
        publish_dt = parse_publish_datetime(str(post.get("publish_time") or ""))
        return base * _calc_time_weight(publish_dt, ref_now)

    def _set_estimated_score_fields(self, post: dict, config: CrawlConfig) -> None:
        total_comments = int(post.get("comments", 0) or 0)
        comment_factor = max(0.5, float(config.topic_comment_factor))
        publish_dt = parse_publish_datetime(str(post.get("publish_time") or ""))
        detail = calculate_score({**post, "author_replies": 0, "publish_dt": publish_dt}, config)

        post["non_author_comments"] = total_comments
        post["author_replies"] = 0
        post["topic_comment_factor"] = comment_factor
        post["base_score"] = detail.base_score
        post["time_weight"] = detail.time_weight
        post["score"] = detail.final_score
        post["score_detail"] = detail.to_dict()
        post.setdefault("top_comment_1", "")
        post.setdefault("top_comment_2", "")
        post.setdefault("top_comment_3", "")
        post.setdefault("top_comment_count", 0)
        post.setdefault("top_comments_data", [])
        post.setdefault("all_comments_data", [])
        post.setdefault("comment_image_urls", "")
        post["comment_analysis_done"] = total_comments <= 0

    def _ensure_candidate_comment_analysis(self, posts: list[dict], config: CrawlConfig) -> bool:
        candidate_pool = _select_weekly_posts(posts, limit=FINAL_COMMENT_ANALYSIS_LIMIT)
        missing = [
            post
            for post in candidate_pool
            if not bool(post.get("comment_analysis_done"))
            and int(post.get("comments", 0) or 0) > 0
            and str(post.get("post_id") or "").strip()
            and str(post.get("author_id") or "").strip()
        ]
        if not missing:
            return False

        total = len(missing)
        worker_count = _bounded_worker_count(config.comment_workers, total)
        self._log(f"补全候选热评与评论结构：{total} 条，{worker_count} 个线程。")
        if worker_count == 1:
            for idx, post in enumerate(missing, start=1):
                self._enrich_score_fields(post, config)
                if idx == total or idx % 5 == 0:
                    self._log(f"候选评论补全进度 {idx}/{total}")
            return True

        with ThreadPoolExecutor(max_workers=worker_count, thread_name_prefix="candidate-comments") as executor:
            futures = {
                executor.submit(self._enrich_score_fields_with_private_session, post, config): post
                for post in missing
            }
            for completed, future in enumerate(as_completed(futures), start=1):
                post = futures[future]
                try:
                    future.result()
                except Exception as err:
                    self._log(
                        f"候选评论补全失败 {completed}/{total}: {post.get('post_id', '-')}, "
                        f"{type(err).__name__}: {err}"
                    )
                if completed == total or completed % 5 == 0:
                    self._log(f"候选评论补全进度 {completed}/{total}")
        return True

    def _enrich_score_fields_with_private_session(self, post: dict, config: CrawlConfig) -> None:
        session = self._new_session()
        try:
            self._enrich_score_fields(post, config, session=session)
        finally:
            session.close()

    def _enrich_score_fields(
        self,
        post: dict,
        config: CrawlConfig,
        session: requests.Session | None = None,
    ) -> None:
        total_comments = int(post.get("comments", 0) or 0)
        author_id = str(post.get("author_id") or "")
        post_id = str(post.get("post_id") or "")

        author_replies = 0
        top_comments: list[dict] = []
        all_comments: list[dict] = []
        analysis_ok = total_comments <= 0
        if total_comments > 0 and post_id and author_id:
            try:
                comment_analysis = self._read_comment_cache(post_id)
                if comment_analysis is None:
                    comment_analysis = self._analyze_comments(
                        post_id=post_id,
                        author_id=author_id,
                        page_limit=min(config.comment_page_limit, max(1, math.ceil(total_comments / 20))),
                        session=session,
                    )
                    self._write_comment_cache(post_id, comment_analysis)
                else:
                    self._log(f"已读取评论缓存：{post_id}")
                author_replies = int(comment_analysis.get("author_replies", 0) or 0)
                top_comments = list(comment_analysis.get("top_comments", []) or [])
                all_comments = list(comment_analysis.get("all_comments", []) or [])
                analysis_ok = True
            except Exception:
                # 评论结构接口不稳定时，不让整批任务失败
                author_replies = 0
                top_comments = []
                all_comments = []

        non_author_comments = max(0, total_comments - author_replies)
        comment_factor = max(0.5, float(config.topic_comment_factor))
        publish_dt = parse_publish_datetime(str(post.get("publish_time") or ""))
        detail = calculate_score({**post, "author_replies": author_replies, "publish_dt": publish_dt}, config)

        post["non_author_comments"] = non_author_comments
        post["author_replies"] = author_replies
        post["topic_comment_factor"] = comment_factor
        post["base_score"] = detail.base_score
        post["time_weight"] = detail.time_weight
        post["score"] = detail.final_score
        post["score_detail"] = detail.to_dict()
        post["top_comment_1"] = _format_comment_for_cell(top_comments[0]) if len(top_comments) >= 1 else ""
        post["top_comment_2"] = _format_comment_for_cell(top_comments[1]) if len(top_comments) >= 2 else ""
        post["top_comment_3"] = _format_comment_for_cell(top_comments[2]) if len(top_comments) >= 3 else ""
        post["top_comment_count"] = len(top_comments)
        post["top_comments_data"] = top_comments
        post["all_comments_data"] = all_comments
        post["comment_image_urls"] = " | ".join(_collect_top_comment_image_urls(top_comments))
        post["comment_analysis_done"] = analysis_ok

    def _read_comment_cache(self, post_id: str) -> dict | None:
        if not self.comment_cache_reader:
            return None
        try:
            data = self.comment_cache_reader(post_id)
        except Exception as err:
            self._log(f"评论缓存无效，重新请求：{post_id}，{type(err).__name__}: {err}")
            return None
        if not isinstance(data, dict):
            return None
        required = {"author_replies", "top_comments", "all_comments"}
        if not required.issubset(data):
            self._log(f"评论缓存无效，重新请求：{post_id}")
            return None
        return data

    def _write_comment_cache(self, post_id: str, data: dict) -> None:
        if not self.comment_cache_writer:
            return
        try:
            self.comment_cache_writer(post_id, data)
        except Exception as err:
            self._log(f"评论缓存写入失败：{post_id}，{type(err).__name__}: {err}")

    def _recalibrate_time_weight(self, posts: list[dict], config: CrawlConfig) -> None:
        rows = list(posts)
        if len(rows) < 12:
            return

        # 只用周报候选帖评估拟合，和实际Top15逻辑保持一致
        row_meta: list[tuple[dict, float, datetime | None, str, float | None]] = []
        all_dist: Counter[str] = Counter()
        ref_now = config.window_end or datetime.now()
        for row in rows:
            publish_time = str(row.get("publish_time", "") or "")
            dt = parse_publish_datetime(publish_time)
            date_key = _date_key_from_publish_time(publish_time, dt)
            base_score = float(row.get("base_score", row.get("score", 0.0)) or 0.0)
            age_ratio = _time_age_ratio(dt, ref_now)
            row_meta.append((row, base_score, dt, date_key, age_ratio))
            all_dist[date_key] += 1

        candidate_meta = [item for item in row_meta if not _should_skip_weekly_post(item[0])]
        candidate_rows = [item[0] for item in candidate_meta]
        if len(candidate_rows) < 12:
            return

        if not all_dist:
            return

        # 扩大强度搜索范围，给近一周这类“沉淀不足窗口”更充分的校准空间
        strengths = [x / 100.0 for x in range(0, 121, 5)]  # 0.00 ~ 1.20
        top_n = min(15, len(candidate_rows))
        target_low, target_high = 0.90, 0.93
        target_mid = (target_low + target_high) / 2.0

        results: list[dict] = []
        for strength in strengths:
            picked = heapq.nlargest(
                top_n,
                (
                    (base_score * _time_weight_from_age_ratio(age_ratio, strength), date_key, row)
                    for row, base_score, _dt, date_key, age_ratio in candidate_meta
                ),
                key=lambda x: x[0],
            )
            picked_dist = Counter(item[1] for item in picked)
            fit_info = _calc_date_distribution_fit(all_dist, picked_dist)
            fit_score = float(fit_info.get("fit_score", 0.0) or 0.0)
            avg_score = sum(item[0] for item in picked) / max(1, top_n)
            in_range = target_low <= fit_score <= target_high
            results.append(
                {
                    "strength": strength,
                    "fit_score": fit_score,
                    "avg_score": avg_score,
                    "in_range": in_range,
                    "fit_gap": abs(fit_score - target_mid),
                }
            )

        if not results:
            return

        best = sorted(
            results,
            key=lambda x: (
                0 if x["in_range"] else 1,
                x["fit_gap"],
                -float(x["avg_score"]),
            ),
        )[0]

        # 兜底：若没有任何候选进入目标区间，则优先取拟合度最高者
        if not any(bool(x["in_range"]) for x in results):
            best = sorted(
                results,
                key=lambda x: (
                    -float(x["fit_score"]),
                    -float(x["avg_score"]),
                ),
            )[0]
            best_fit = max(float(x["fit_score"]) for x in results)
            self._log(
                "注意：当前窗口仅靠时间权重难以达到目标拟合区间，"
                f"本次可达到的最高预估拟合度为 {best_fit * 100:.2f}%"
            )

        chosen_strength = float(best["strength"])
        for row, base_score, _dt, _date_key, age_ratio in row_meta:
            w = _time_weight_from_age_ratio(age_ratio, chosen_strength)
            row["time_weight"] = round(w, 4)
            row["score"] = round(base_score * w, 4)

        self._log(
            "时间权重校准完成："
            f"strength={chosen_strength:.2f}，"
            f"预估拟合度={float(best['fit_score']) * 100:.2f}%"
        )

    def _analyze_comments(
        self,
        post_id: str,
        author_id: str,
        page_limit: int,
        session: requests.Session | None = None,
    ) -> dict:
        http = session or self.session
        max_id = "0"
        seen_root_ids: set[str] = set()
        seen_comment_ids: set[str] = set()
        root_reply_counts: dict[str, int] = {}
        comment_candidates: list[dict] = []

        for _ in range(max(1, page_limit)):
            params = {
                "id": post_id,
                "is_reload": "1",
                "is_show_bulletin": "2",
                "is_mix": "0",
                "flow": "1",
                "count": "20",
            }
            if max_id != "0":
                params["max_id"] = max_id

            resp = http.get(
                COMMENTS_API_URL,
                params=params,
                headers={"Referer": f"https://weibo.com/detail/{post_id}", "Origin": "https://weibo.com"},
                timeout=20,
            )
            if resp.status_code >= 400:
                break

            payload = resp.json()
            if int(payload.get("ok", 0) or 0) != 1:
                break

            data = payload.get("data") or []
            if not isinstance(data, list) or not data:
                break

            for root in data:
                root_comment_id = str(root.get("id") or "")
                if root_comment_id and root_comment_id not in seen_comment_ids:
                    seen_comment_ids.add(root_comment_id)
                    comment_candidates.append(
                        {
                            "user_name": str(((root.get("user") or {}).get("screen_name")) or ""),
                            "like_counts": int(root.get("like_counts") or 0),
                            "created_at": str(root.get("created_at") or ""),
                            "text": str(root.get("text_raw") or root.get("text") or ""),
                            "image_urls": _extract_comment_image_urls(root),
                        }
                    )

                root_id = str(root.get("id") or root.get("rootid") or "")
                if not root_id:
                    continue
                if root_id in seen_root_ids:
                    continue
                seen_root_ids.add(root_id)

                replies = root.get("comments") or []
                if not isinstance(replies, list):
                    continue
                for reply in replies:
                    reply_id = str(reply.get("id") or "")
                    if reply_id and reply_id not in seen_comment_ids:
                        seen_comment_ids.add(reply_id)
                        comment_candidates.append(
                            {
                                "user_name": str(((reply.get("user") or {}).get("screen_name")) or ""),
                                "like_counts": int(reply.get("like_counts") or 0),
                                "created_at": str(reply.get("created_at") or ""),
                                "text": str(reply.get("text_raw") or reply.get("text") or ""),
                                "image_urls": _extract_comment_image_urls(reply),
                            }
                        )

                    uid = str(((reply.get("user") or {}).get("id")) or "")
                    if uid != author_id:
                        continue
                    thread_id = str(reply.get("rootid") or root_id)
                    root_reply_counts[thread_id] = min(root_reply_counts.get(thread_id, 0) + 1, 3)

            max_id = str(payload.get("max_id") or "0")
            if max_id == "0":
                break

        top_comments = heapq.nlargest(
            3,
            comment_candidates,
            key=lambda x: (int(x.get("like_counts", 0) or 0), str(x.get("created_at", ""))),
        )

        return {
            "author_replies": sum(root_reply_counts.values()),
            "top_comments": top_comments,
            "all_comments": comment_candidates,
        }

    def hydrate_full_text_posts(self, posts: list[dict], max_workers: int | None = None) -> None:
        all_rows = list(posts)
        rows = [post for post in all_rows if _needs_full_text_hydration(post)]
        total = len(rows)
        skipped = len(all_rows) - total
        if skipped > 0:
            self._log(f"正文校正筛选：{total}/{len(all_rows)} 条需要网络补全，跳过 {skipped} 条完整正文。")
        if not total:
            self._log("正文校正：未发现疑似截断正文，跳过网络补全。")
            return

        worker_count = _bounded_worker_count(max_workers or 6, total)
        if worker_count == 1:
            for idx, post in enumerate(rows, start=1):
                post_id = str(post.get("post_id") or "")
                self._log(f"正文校正 {idx}/{total}: {post_id or '-'}")
                self._hydrate_one_post(post)
            return

        self._log(f"正文校正并行处理：{worker_count} 个线程。")
        with ThreadPoolExecutor(max_workers=worker_count, thread_name_prefix="text-hydrate") as executor:
            futures = {
                executor.submit(self._hydrate_one_post_with_private_session, post): post
                for post in rows
            }
            for completed, future in enumerate(as_completed(futures), start=1):
                post = futures[future]
                try:
                    future.result()
                except Exception as err:
                    self._log(
                        f"正文校正失败 {completed}/{total}: {post.get('post_id', '-')}, "
                        f"{type(err).__name__}: {err}"
                    )
                if completed == total or completed % 5 == 0:
                    self._log(f"正文校正进度 {completed}/{total}")

    def _hydrate_one_post_with_private_session(self, post: dict) -> None:
        session = self._new_session()
        try:
            self._hydrate_one_post(post, session=session)
        finally:
            session.close()

    def _hydrate_one_post(self, post: dict, session: requests.Session | None = None) -> None:
        content = str(post.get("content") or "")
        post_id = str(post.get("post_id") or "")
        if not post_id and not str(post.get("post_url") or "").strip():
            return
        full_text, extra_image_urls = self._fetch_post_full_text(post, session=session)
        cleaned_old = _strip_specific_topic_tags(_remove_expand_hint(_clean_text(content)))
        report_old = _strip_specific_topic_tags_preserve(_remove_expand_hint_preserve(content))
        cleaned_new = _strip_specific_topic_tags(_remove_expand_hint(_clean_text(full_text)))
        report_new = _strip_specific_topic_tags_preserve(_remove_expand_hint_preserve(full_text))
        if _should_replace_content(cleaned_old, cleaned_new):
            post["content"] = report_new or cleaned_new
        else:
            post["content"] = report_old or cleaned_old
        if extra_image_urls:
            merged = _dedup_keep_order(
                _split_multi_urls(str(post.get("original_image_urls") or ""), sep="|") + extra_image_urls
            )
            post["original_image_urls"] = " | ".join(merged)
            post["image_count"] = len(merged)

    def _fetch_post_full_text(
        self,
        post: dict,
        session: requests.Session | None = None,
    ) -> tuple[str, list[str]]:
        http = session or self.session
        post_id = str(post.get("post_id") or "")
        post_url = str(post.get("post_url") or "")
        referer = post_url or (f"https://weibo.com/detail/{post_id}" if post_id else "https://weibo.com/")
        full_text = ""
        image_urls: list[str] = []

        if post_id:
            try:
                resp = http.get(
                    "https://weibo.com/ajax/statuses/show",
                    params={"id": post_id},
                    headers={"Referer": referer, "Origin": "https://weibo.com"},
                    timeout=20,
                )
                if resp.status_code == 200:
                    data = resp.json() if resp.content else {}
                    if isinstance(data, dict):
                        raw_text = data.get("text_raw") or data.get("longTextContent_raw") or ""
                        if not raw_text:
                            raw_text = _html_to_text_preserve(str(data.get("text") or data.get("longTextContent") or ""))
                        full_text = _clean_text_preserve(str(raw_text or ""))
                        image_urls = _extract_status_image_urls(data)
            except Exception:
                pass

        if post_id and _looks_truncated_text(full_text):
            try:
                long_resp = http.get(
                    "https://weibo.com/ajax/statuses/longtext",
                    params={"id": post_id},
                    headers={"Referer": referer, "Origin": "https://weibo.com"},
                    timeout=20,
                )
                if long_resp.status_code == 200:
                    payload = long_resp.json() if long_resp.content else {}
                    if isinstance(payload, dict):
                        data = payload.get("data") if isinstance(payload.get("data"), dict) else payload
                        raw_text = str((data or {}).get("longTextContent_raw") or "")
                        if not raw_text:
                            raw_text = _html_to_text_preserve(str((data or {}).get("longTextContent") or ""))
                        raw_text = _clean_text_preserve(raw_text)
                        if raw_text and _should_replace_content(full_text, raw_text):
                            full_text = raw_text
                        image_urls = _dedup_keep_order(image_urls + _extract_status_image_urls(data or {}))
            except Exception:
                pass

        if post_id and _looks_truncated_text(full_text):
            try:
                mobile_resp = http.get(
                    "https://m.weibo.cn/statuses/extend",
                    params={"id": post_id},
                    headers={"Referer": referer},
                    timeout=20,
                )
                if mobile_resp.status_code == 200 and mobile_resp.content:
                    payload = mobile_resp.json()
                    data = payload.get("data") if isinstance(payload, dict) else {}
                    mobile_text = _html_to_text_preserve(str((data or {}).get("longTextContent") or ""))
                    mobile_text = _clean_text_preserve(mobile_text)
                    if mobile_text and _should_replace_content(full_text, mobile_text):
                        full_text = mobile_text
            except Exception:
                pass

        if (not full_text or _looks_truncated_text(full_text)) and post_url:
            try:
                page = http.get(post_url, headers={"Referer": "https://weibo.com/"}, timeout=20)
                if page.status_code == 200 and page.text:
                    try:
                        html = extract_feed_html_from_page(page.text)
                        parsed = parse_posts_from_html(html)
                    except Exception:
                        parsed = []
                    if parsed:
                        matched = next((x for x in parsed if str(x.get("post_id") or "") == post_id), parsed[0])
                        full_text = _clean_text_preserve(str(matched.get("content") or full_text))
                        image_urls = _dedup_keep_order(
                            image_urls + _split_multi_urls(str(matched.get("original_image_urls") or ""), sep="|")
                        )
                    else:
                        detail_text = _extract_full_text_from_detail_html(page.text)
                        if detail_text:
                            full_text = _clean_text_preserve(detail_text)
                        elif not full_text:
                            blob_text = _extract_text_raw_from_page_blob(page.text)
                            if blob_text:
                                full_text = _clean_text_preserve(blob_text)
            except Exception:
                pass

        return _remove_expand_hint_preserve(full_text), image_urls


def parse_super_topic_id(input_text: str) -> str | None:
    return _parse_super_topic_id(input_text)


def build_report_title(topic_name: str | None = None, super_topic: str | None = None) -> str:
    name = normalize_super_topic_name(topic_name or "")
    if not name:
        name = normalize_super_topic_name(str(super_topic or ""))
    return f"{name or '微博'}超话周报"


def normalize_super_topic_name(value: str) -> str:
    raw = _clean_text(str(value or ""))
    if not raw:
        return ""
    raw = raw.strip().strip("#")
    raw = re.sub(r"^https?://\S+$", "", raw, flags=re.I)
    raw = re.sub(r"^100808[0-9a-fA-F]+$", "", raw)
    raw = re.sub(r"\s*[-_｜|].*$", "", raw)
    raw = re.sub(r"(?:微博)?超话(?:社区|详情|首页|主页)?$", "", raw)
    raw = re.sub(r"(?:的)?微博(?:主页)?$", "", raw)
    raw = raw.strip(" #　-—_｜|：:")
    if not raw or raw.lower() in {"weibo", "m.weibo.cn", "weibo.com"}:
        return ""
    return raw[:40]


def extract_super_topic_name(page_html: str, fallback: str | None = None) -> str:
    html = str(page_html or "")
    candidates: list[str] = []

    title_match = re.search(r"<title[^>]*>(.*?)</title>", html, flags=re.I | re.S)
    if title_match:
        candidates.append(_html_to_text(title_match.group(1)))

    for pattern in (
        r'<meta[^>]+(?:property|name)=["\'](?:og:title|keywords|description)["\'][^>]+content=["\']([^"\']+)["\']',
        r'<meta[^>]+content=["\']([^"\']+)["\'][^>]+(?:property|name)=["\'](?:og:title|keywords|description)["\']',
    ):
        candidates.extend(_html_to_text(match.group(1)) for match in re.finditer(pattern, html, flags=re.I | re.S))

    topic_match = re.search(r"#?\s*([^#<>{}\"'，,。；;｜|]{1,40}?)\s*超话", html)
    if topic_match:
        candidates.append(topic_match.group(1))

    candidates.append(str(fallback or ""))

    for candidate in candidates:
        name = normalize_super_topic_name(candidate)
        if name:
            return name
    return ""


def extract_feed_html_from_page(page_html: str) -> str:
    objects = parse_fm_view_objects(page_html)
    if not objects:
        raise CrawlError("页面中未找到 FM.view 数据块，无法解析帖子列表。")

    for obj in objects:
        domid = str(obj.get("domid", ""))
        html = obj.get("html")
        if (
            domid.startswith("Pl_Core_MixedFeed__")
            and isinstance(html, str)
            and "feed_list_item" in html
        ):
            return html

    html_candidates = [
        obj.get("html", "")
        for obj in objects
        if isinstance(obj.get("html"), str) and "feed_list_item" in obj.get("html", "")
    ]
    if html_candidates:
        return max(html_candidates, key=len)

    raise CrawlError("页面结构已变化：未在 FM.view 中找到帖子列表 HTML。")


def parse_fm_view_objects(page_html: str) -> list[dict]:
    results: list[dict] = []
    cursor = 0
    while True:
        idx = page_html.find(FM_VIEW_MARKER, cursor)
        if idx < 0:
            break

        payload_start = idx + len(FM_VIEW_MARKER)
        payload_start = _skip_space(page_html, payload_start)
        if payload_start >= len(page_html) or page_html[payload_start] != "{":
            cursor = payload_start + 1
            continue

        payload_end = _find_json_object_end(page_html, payload_start)
        if payload_end < 0:
            cursor = payload_start + 1
            continue

        payload = page_html[payload_start : payload_end + 1]
        try:
            obj = json.loads(payload)
            if isinstance(obj, dict):
                results.append(obj)
        except Exception:
            pass
        cursor = payload_end + 1
    return results


def _skip_space(text: str, start: int) -> int:
    i = start
    while i < len(text) and text[i] in (" ", "\n", "\r", "\t"):
        i += 1
    return i


def _find_json_object_end(text: str, obj_start: int) -> int:
    depth = 0
    in_string = False
    escaped = False
    for i, ch in enumerate(text[obj_start:], start=obj_start):
        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            continue
        if ch == '"':
            in_string = True
        elif ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return i
    return -1


def parse_posts_from_html(html: str) -> list[dict]:
    try:
        soup = BeautifulSoup(html, "lxml")
    except Exception:
        soup = BeautifulSoup(html, "html.parser")
    items = [item for item in soup.select("div[action-type='feed_list_item']") if not _is_nested_feed_item(item)]
    posts: list[dict] = []
    for item in items:
        post_id = item.get("mid") or item.get("data-mid") or ""
        user_name = _extract_user_name(item)
        author_id = _extract_author_id(item)
        publish_time = _extract_publish_time(item)
        post_url = _extract_post_url(item)
        original_image_urls = _extract_original_image_urls(item)
        content = _strip_specific_topic_tags_preserve(_remove_expand_hint_preserve(_extract_content(item)))
        has_video = _extract_has_video(item)
        reposts = _extract_action_count(item, ["feed_list_forward", "fl_forward"])
        comments = _extract_action_count(item, ["feed_list_comment", "fl_comment"])
        likes = _extract_action_count(item, ["feed_list_like", "fl_like"])

        posts.append(
            {
                "post_id": str(post_id),
                "author_id": str(author_id),
                "user_name": user_name,
                "publish_time": publish_time,
                "post_url": post_url,
                "original_image_urls": " | ".join(original_image_urls),
                "image_count": len(original_image_urls),
                "downloaded_image_count": 0,
                "image_local_paths": "",
                "content": content,
                "has_video": has_video,
                "reposts": reposts,
                "comments": comments,
                "likes": likes,
                "non_author_comments": 0,
                "author_replies": 0,
                "topic_comment_factor": 1.0,
                "score": 0.0,
                "top_comment_1": "",
                "top_comment_2": "",
                "top_comment_3": "",
                "top_comment_count": 0,
                "engagement_total": reposts + comments + likes,
            }
        )
    return posts


def _extract_full_text_from_detail_html(page_html: str) -> str:
    try:
        soup = BeautifulSoup(page_html, "lxml")
    except Exception:
        soup = BeautifulSoup(page_html, "html.parser")
    selectors = [
        "div[node-type='feed_list_content_full']",
        "p[node-type='feed_list_content_full']",
        "div.WB_text",
        "p.WB_text",
    ]
    texts: list[str] = []
    for selector in selectors:
        for node in soup.select(selector):
            if _is_inside_forwarded_content(node):
                continue
            text = _clean_text_preserve(node.get_text("\n", strip=True))
            if text and "微博社区公约" not in text:
                texts.append(text)
    if not texts:
        return ""
    return max(texts, key=len)


def _extract_text_raw_from_page_blob(page_html: str) -> str:
    # 兜底：从详情页脚本中的 JSON 片段尝试抓 text_raw，避免“展开全文”残留。
    m = re.search(r'"text_raw"\s*:\s*"((?:\\.|[^"\\])*)"', page_html)
    if not m:
        return ""
    raw = m.group(1)
    try:
        decoded = json.loads(f"\"{raw}\"")
    except Exception:
        return ""
    return _html_to_text_preserve(str(decoded or ""))


def parse_publish_datetime(text: str) -> datetime | None:
    return parse_weibo_time(text)


def _parse_publish_datetime_with_format(raw: str, fmt: str) -> datetime | None:
    try:
        return datetime.strptime(raw, fmt)
    except ValueError:
        return None


def _apply_carryover_bucket(publish_dt: datetime | None, carryover_hours: int) -> datetime | None:
    """把截止前若干小时的帖子顺延到下一期（通过归档时间平移实现）。"""
    if publish_dt is None:
        return None
    hours = max(0, int(carryover_hours or 0))
    if hours <= 0:
        return publish_dt
    return publish_dt + timedelta(hours=hours)


def _bounded_worker_count(requested: int | None, total: int) -> int:
    if total <= 0:
        return 1
    try:
        desired = int(requested or 1)
    except (TypeError, ValueError):
        desired = 1
    return max(1, min(desired, total, 12))


def _calc_time_weight(publish_dt: datetime | None, now: datetime, strength: float = 0.06) -> float:
    """轻量时间权重：strength 越大，时间偏置越强。"""
    return _time_weight_from_age_ratio(_time_age_ratio(publish_dt, now), strength)


def _time_age_ratio(publish_dt: datetime | None, now: datetime) -> float | None:
    if publish_dt is None:
        return None
    age_hours = max(0.0, (now - publish_dt).total_seconds() / 3600.0)
    return min(1.0, age_hours / (7.0 * 24.0))


def _time_weight_from_age_ratio(age_ratio: float | None, strength: float = 0.06) -> float:
    if age_ratio is None:
        return 1.0
    # 中心约 1.01；strength=0.06 时范围约 [1.04, 0.98]
    # 为避免极端强度时旧帖权重过低，设置下限保护。
    s = max(0.0, float(strength))
    weight = 1.01 + s * (0.5 - age_ratio)
    return max(0.75, weight)


def _should_replace_content(old_text: str, new_text: str) -> bool:
    if not new_text:
        return False
    if not old_text:
        return True
    if old_text == new_text:
        return False
    if _looks_truncated_text(old_text):
        return True
    # 新内容明显更完整时替换
    if len(new_text) >= len(old_text) + 18:
        return True
    # 旧内容疑似截断，且新内容稍长时替换
    return _looks_truncated_text(old_text) and len(new_text) >= len(old_text) + 6


def _looks_truncated_text(text: str) -> bool:
    raw = _clean_text(text)
    if not raw:
        return False
    if "展开全文" in raw or "展开原文" in raw:
        return True
    if raw.endswith(("...", "…", "⋯")):
        return True
    # 超长但以半截字符结尾，通常是列表页截断文本
    return len(raw) >= 140 and not re.search(r"[。！？!?；;）)\]】”\"]$", raw)


def _needs_full_text_hydration(post: dict) -> bool:
    content = str(post.get("content") or "")
    cleaned = _clean_text(content)
    if not cleaned:
        return True
    if _looks_truncated_text(cleaned):
        return True
    return "展开全文" in content or "展开原文" in content


def _remove_expand_hint(text: str) -> str:
    raw = _clean_text(text)
    if not raw:
        return ""
    raw = re.sub(r"(展开全文|展开原文)\s*[cC]?", " ", raw)
    raw = re.sub(r"\s{2,}", " ", raw)
    return raw.strip()


def _remove_expand_hint_preserve(text: str) -> str:
    raw = _clean_text_preserve(text)
    if not raw:
        return ""
    raw = re.sub(r"(展开全文|展开原文)\s*[cC]?", " ", raw)
    return _clean_text_preserve(raw)


def export_posts_csv(posts: Iterable[dict], csv_path: Path) -> None:
    _export_posts_csv(posts, csv_path, EXPORT_COLUMN_MAP)


def export_posts_xlsx(posts: Iterable[dict], xlsx_path: Path) -> None:
    xlsx_path.parent.mkdir(parents=True, exist_ok=True)
    rows = list(posts)
    wb = Workbook()
    ws = cast(Worksheet, wb.active)
    ws.title = "帖子统计"

    base_headers = [cn for _, cn in EXPORT_COLUMN_MAP]
    max_embed_images = max(
        (len(_get_embed_image_paths(row)) for row in rows),
        default=0,
    )
    image_headers = [f"{EMBED_IMAGE_HEADER_PREFIX}{i}" for i in range(1, max_embed_images + 1)]
    headers = base_headers + image_headers
    ws.append(headers)

    header_font = Font(bold=True)
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        if header.startswith(EMBED_IMAGE_HEADER_PREFIX):
            width = 26
        else:
            width = WIDE_COLUMNS.get(header, 14)
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    for post in rows:
        row = _build_export_row(post)
        ws.append([row[h] for h in base_headers] + [""] * max_embed_images)

    wrap_cols = {"帖子内容", "原图链接", "图片本地路径", "热评1(按点赞)", "热评2(按点赞)", "热评3(按点赞)"}
    for col_idx, header in enumerate(headers, start=1):
        if header in wrap_cols:
            for r in range(2, ws.max_row + 1):
                ws.cell(row=r, column=col_idx).alignment = Alignment(wrap_text=True, vertical="top")

    if max_embed_images > 0:
        base_col_count = len(base_headers)
        for row_idx, post in enumerate(rows, start=2):
            image_paths = _get_embed_image_paths(post)
            if not image_paths:
                continue
            ws.row_dimensions[row_idx].height = 125
            for img_idx, img_path in enumerate(image_paths[:max_embed_images], start=1):
                col_idx = base_col_count + img_idx
                cell_ref = f"{get_column_letter(col_idx)}{row_idx}"
                try:
                    xl_img = XLImage(img_path)
                except Exception:
                    continue
                max_w, max_h = 180, 120
                if xl_img.width and xl_img.height:
                    scale = min(max_w / xl_img.width, max_h / xl_img.height, 1.0)
                    xl_img.width = int(xl_img.width * scale)
                    xl_img.height = int(xl_img.height * scale)
                ws.add_image(xl_img, cell_ref)

    wb.save(xlsx_path)


def _build_export_row(post: dict) -> dict:
    return _export_build_row(post, EXPORT_COLUMN_MAP)


def build_comment_leaderboards(posts: Iterable[dict], top_n: int = 3) -> dict:
    rows = list(posts)
    stats: dict[str, dict] = {}

    def ensure_user(name: str) -> dict:
        key = _clean_text(name) or "匿名用户"
        if key not in stats:
            stats[key] = {
                "user_name": key,
                "comment_count": 0,
                "commented_post_count": 0,
                "comment_likes_total": 0,
                "hot_top3_count": 0,
                "like_rate": 0.0,
                "hot_rate": 0.0,
                "quality_score": 0.0,
                "_commented_post_ids": set(),
            }
        return stats[key]

    for post in rows:
        post_id = _clean_text(str(post.get("post_id", "") or ""))
        post_key = post_id or f"{_clean_text(str(post.get('post_url', '') or ''))}|{_clean_text(str(post.get('publish_time', '') or ''))}"
        all_comments = list(post.get("all_comments_data") or [])
        for comment in all_comments:
            if not isinstance(comment, dict):
                continue
            item = ensure_user(str(comment.get("user_name", "") or "匿名用户"))
            item["comment_count"] += 1
            item["comment_likes_total"] += int(comment.get("like_counts", 0) or 0)
            item["_commented_post_ids"].add(post_key)

        hot_comments = list(post.get("top_comments_data") or [])[:3]
        for comment in hot_comments:
            if not isinstance(comment, dict):
                continue
            item = ensure_user(str(comment.get("user_name", "") or "匿名用户"))
            item["hot_top3_count"] += 1

    users = [x for x in stats.values() if int(x.get("comment_count", 0) or 0) > 0]
    if not users:
        return {"comment_count_top3": [], "comment_quality_top3": [], "all_stats": []}

    for item in users:
        count = max(1, int(item["comment_count"]))
        item["like_rate"] = float(item["comment_likes_total"]) / count
        item["hot_rate"] = float(item["hot_top3_count"]) / count
        item["commented_post_count"] = len(item.get("_commented_post_ids", set()))

    max_like_rate = max((float(x["like_rate"]) for x in users), default=0.0)
    max_hot_rate = max((float(x["hot_rate"]) for x in users), default=0.0)

    for item in users:
        like_norm = (float(item["like_rate"]) / max_like_rate) if max_like_rate > 0 else 0.0
        hot_norm = (float(item["hot_rate"]) / max_hot_rate) if max_hot_rate > 0 else 0.0
        base_score = 0.6 * like_norm + 0.4 * hot_norm
        stability = min(1.0, float(item["comment_count"]) / 8.0)
        item["quality_score"] = round(base_score * stability, 4)

    count_sorted = sorted(
        users,
        key=lambda x: (
            int(x["comment_count"]),
            int(x.get("commented_post_count", 0)),
            int(x["comment_likes_total"]),
            int(x["hot_top3_count"]),
            str(x["user_name"]),
        ),
        reverse=True,
    )
    quality_sorted = sorted(
        users,
        key=lambda x: (
            float(x["quality_score"]),
            int(x["comment_likes_total"]),
            int(x["hot_top3_count"]),
            int(x["comment_count"]),
            str(x["user_name"]),
        ),
        reverse=True,
    )

    def with_rank(rows_: list[dict]) -> list[dict]:
        out: list[dict] = []
        for idx, item in enumerate(rows_[: max(1, top_n)], start=1):
            copied = dict(item)
            copied.pop("_commented_post_ids", None)
            copied["rank"] = idx
            out.append(copied)
        return out

    all_stats: list[dict] = []
    for item in users:
        copied = dict(item)
        copied.pop("_commented_post_ids", None)
        all_stats.append(copied)

    return {
        "comment_count_top3": with_rank(count_sorted),
        "comment_quality_top3": with_rank(quality_sorted),
        "all_stats": all_stats,
    }


def _format_leaderboard_line(
    item: dict,
    include_hot: bool = True,
    include_quality: bool = False,
    include_like_total: bool = True,
    include_post_span: bool = False,
) -> str:
    rank = int(item.get("rank", 0) or 0)
    user_name = _clean_text(str(item.get("user_name", "") or "匿名用户"))
    comment_count = int(item.get("comment_count", 0) or 0)
    commented_post_count = int(item.get("commented_post_count", 0) or 0)
    like_total = int(item.get("comment_likes_total", 0) or 0)
    hot_count = int(item.get("hot_top3_count", 0) or 0)
    line = f"{rank}. @{user_name}：评论 {comment_count} 条"
    if include_post_span:
        line += f"，评论过 {commented_post_count} 条帖子"
    elif include_like_total:
        line += f"，本周评论获赞 {like_total}"
    if include_hot:
        line += f"，热评前三 {hot_count} 次"
    if include_quality:
        line += f"，质量分 {float(item.get('quality_score', 0.0)):.4f}"
    return line


def export_weekly_report_docx(
    posts: Iterable[dict],
    docx_path: Path,
    title: str = "微博超话周报",
    leaderboards: dict | None = None,
    preselected: bool = False,
    max_bytes: int = DOCX_SIZE_LIMIT_BYTES,
) -> list[Path]:
    all_posts = list(posts)
    rows = all_posts[:15] if preselected else _select_weekly_posts(all_posts, limit=15)
    board = leaderboards or build_comment_leaderboards(all_posts, top_n=3)
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    if not rows:
        out_path = _numbered_docx_path(docx_path, 1)
        _write_weekly_report_docx_part([], out_path, title, board, include_leaderboards=True)
        return [out_path]

    parts: list[Path] = []
    current: list[dict] = []
    trial_path = docx_path.with_name(f"_{docx_path.stem}_trial{docx_path.suffix}")
    limit = max(1, int(max_bytes or DOCX_SIZE_LIMIT_BYTES))

    for post in rows:
        trial_rows = current + [post]
        _write_weekly_report_docx_part(
            trial_rows,
            trial_path,
            title,
            board,
            include_leaderboards=(len(parts) == 0),
        )
        trial_size = trial_path.stat().st_size if trial_path.exists() else 0
        with suppress(Exception):
            trial_path.unlink(missing_ok=True)

        if trial_size > limit and current:
            out_path = _numbered_docx_path(docx_path, len(parts) + 1)
            _write_weekly_report_docx_part(
                current,
                out_path,
                title,
                board,
                include_leaderboards=(len(parts) == 0),
            )
            parts.append(out_path)
            current = [post]
        else:
            current = trial_rows

    if current:
        out_path = _numbered_docx_path(docx_path, len(parts) + 1)
        _write_weekly_report_docx_part(
            current,
        out_path,
        title,
        board,
        include_leaderboards=(len(parts) == 0),
    )
        parts.append(out_path)

    return parts


def export_weekly_report_sum_docx(
    posts: Iterable[dict],
    docx_path: Path,
    title: str = "微博超话周报",
    leaderboards: dict | None = None,
    preselected: bool = False,
) -> Path:
    all_posts = list(posts)
    rows = all_posts[:15] if preselected else _select_weekly_posts(all_posts, limit=15)
    board = leaderboards or build_comment_leaderboards(all_posts, top_n=3)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    _write_weekly_report_docx_part(
        rows,
        docx_path,
        title,
        board,
        include_leaderboards=True,
    )
    return docx_path


def _numbered_docx_path(docx_path: Path, seq: int) -> Path:
    return docx_path.with_name(f"{docx_path.stem}_{seq:02d}{docx_path.suffix}")


def _write_weekly_report_docx_part(
    rows: list[dict],
    docx_path: Path,
    title: str,
    board: dict,
    include_leaderboards: bool,
) -> None:
    doc = create_doc()

    normal = cast(Any, doc.styles["Normal"])
    normal.font.name = "Microsoft YaHei"
    _set_east_asia_font(normal, "Microsoft YaHei")
    normal.font.size = Pt(11)

    if include_leaderboards:
        date_range_text = _format_posts_date_range(rows)

        p_title = doc.add_paragraph()
        p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p_title.add_run(title)
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(33, 37, 41)
        run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run, "Microsoft YaHei")

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p_sub.add_run(f"帖子选取日期：{date_range_text}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(52, 58, 64)
        run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run, "Microsoft YaHei")

        p_rank_title = doc.add_paragraph()
        run = p_rank_title.add_run("本周社区互动榜")
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run, "Microsoft YaHei")

        p_count_title = doc.add_paragraph()
        run = p_count_title.add_run("评论数量榜 Top3")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run, "Microsoft YaHei")
        count_rows = list(board.get("comment_count_top3") or [])
        if count_rows:
            for item in count_rows:
                p = doc.add_paragraph()
                run = p.add_run(
                    _format_leaderboard_line(
                        item,
                        include_hot=False,
                        include_quality=False,
                        include_like_total=False,
                        include_post_span=True,
                    )
                )
                run.font.size = Pt(10)
                run.font.name = "Microsoft YaHei"
                _set_run_east_asia_font(run, "Microsoft YaHei")
        else:
            doc.add_paragraph("暂无评论数据")

        p_quality_title = doc.add_paragraph()
        run = p_quality_title.add_run("评论质量榜 Top3")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run, "Microsoft YaHei")
        quality_rows = list(board.get("comment_quality_top3") or [])
        if quality_rows:
            for item in quality_rows:
                p = doc.add_paragraph()
                run = p.add_run(_format_leaderboard_line(item, include_hot=True, include_quality=False))
                run.font.size = Pt(10)
                run.font.name = "Microsoft YaHei"
                _set_run_east_asia_font(run, "Microsoft YaHei")
        else:
            doc.add_paragraph("暂无评论数据")

        doc.add_paragraph("")

        p_posts_title = doc.add_paragraph()
        run = p_posts_title.add_run("本周热帖Top15")
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run, "Microsoft YaHei")
        doc.add_paragraph("")

    for post in rows:
        author = _clean_text(str(post.get("user_name", "") or "未知作者"))
        content = _clean_report_text(str(post.get("content", "") or ""))
        publish_time = _clean_text(str(post.get("publish_time", "") or ""))
        post_url = _clean_text(str(post.get("post_url", "") or ""))

        p_body = doc.add_paragraph()
        run = p_body.add_run(f"@{author}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run, "Microsoft YaHei")
        run_colon = p_body.add_run("：")
        run_colon.bold = True
        run_colon.font.size = Pt(12)
        run_colon.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run_colon, "Microsoft YaHei")
        run2 = p_body.add_run()
        run2.font.size = Pt(12)
        run2.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run2, "Microsoft YaHei")
        _add_preserved_text(run2, content)

        p_time = doc.add_paragraph()
        run = p_time.add_run(f"发送时间：{publish_time}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(73, 80, 87)
        run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run, "Microsoft YaHei")

        image_paths = _split_multi_urls(str(post.get("image_local_paths") or ""), sep="|")
        for img_path in image_paths:
            _add_scaled_right_aligned_picture(doc, img_path, scale=0.50)

        comments = _iter_report_comments(post)
        if comments:
            p_c_title = doc.add_paragraph()
            run = p_c_title.add_run("热评：")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(90, 98, 104)
            run.font.name = "Microsoft YaHei"
            _set_run_east_asia_font(run, "Microsoft YaHei")

            for c in comments:
                c_simple = _clean_report_text(_format_hot_comment_text(c))
                if c_simple:
                    p_c = doc.add_paragraph()
                    run = p_c.add_run()
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(108, 117, 125)
                    run.font.name = "Microsoft YaHei"
                    _set_run_east_asia_font(run, "Microsoft YaHei")
                    _add_preserved_text(run, c_simple)

                for c_img in _split_multi_urls(str(c.get("image_local_paths") or ""), sep="|"):
                    _add_scaled_right_aligned_picture(doc, c_img, scale=0.25)

        p_link = doc.add_paragraph()
        label_run = p_link.add_run("帖子链接：")
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = RGBColor(0, 102, 204)
        label_run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(label_run, "Microsoft YaHei")
        run = _add_external_hyperlink_run(p_link, post_url, post_url) if post_url else p_link.add_run("")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.name = "Microsoft YaHei"
        _set_run_east_asia_font(run, "Microsoft YaHei")

        doc.add_paragraph("")

    doc.save(str(docx_path))


def _add_preserved_text(run, text: str) -> None:
    parts = _clean_text_preserve(text).split("\n")
    for idx, part in enumerate(parts):
        if idx > 0:
            run.add_break()
        if part:
            run.add_text(part)


def _add_external_hyperlink_run(paragraph, text: str, url: str):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.underline = True
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)
    return run


def _add_scaled_right_aligned_picture(doc: DocxDocument, image_path: str, scale: float) -> None:
    try:
        section = doc.sections[-1]
        page_width = int(section.page_width or 0)
        left_margin = int(section.left_margin or 0)
        right_margin = int(section.right_margin or 0)
        usable_width = page_width - left_margin - right_margin
        width = int(usable_width * max(0.05, min(1.0, float(scale))))
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.add_run().add_picture(image_path, width=width)
    except Exception:
        return


def export_weekly_report_md(
    posts: Iterable[dict],
    md_path: Path,
    title: str = "微博超话周报",
    leaderboards: dict | None = None,
    preselected: bool = False,
) -> None:
    _export_weekly_report_md(
        posts,
        md_path,
        title=title,
        leaderboards=leaderboards,
        preselected=preselected,
    )


def download_post_images(
    posts: Iterable[dict],
    image_dir: Path,
    cookie: str,
    user_agent: str | None = None,
    progress_callback: Callable[[str], None] | None = None,
    max_workers: int | None = None,
    cancel_checker: Callable[[], None] | None = None,
) -> None:
    image_dir.mkdir(parents=True, exist_ok=True)
    ua = user_agent or (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
    headers = {
        "User-Agent": ua,
        "Cookie": cookie.strip(),
        "Referer": "https://weibo.com/",
    }

    def _new_session() -> requests.Session:
        session = requests.Session()
        session.headers.update(headers)
        return session

    def _download_url_list(
        session: requests.Session,
        urls: list[str],
        target_dir: Path,
        stem: str,
        prefix: str,
    ) -> list[str]:
        target_dir.mkdir(parents=True, exist_ok=True)
        local_paths: list[str] = []
        uniq_urls = _dedup_image_urls(urls)
        for i, url in enumerate(uniq_urls, start=1):
            if cancel_checker:
                cancel_checker()
            ext = _guess_image_ext(url)
            digest = hashlib.md5(url.encode("utf-8")).hexdigest()[:10]
            filename = f"{stem}_{prefix}_{i}_{digest}{ext}"
            path = target_dir / filename

            if not path.exists():
                try:
                    resp = session.get(url, timeout=20)
                    if resp.status_code == 200 and resp.content:
                        path.write_bytes(resp.content)
                except Exception:
                    continue
            if path.exists():
                local_paths.append(str(path.resolve()))
        return local_paths

    def _download_one_post(idx: int, post: dict, total: int) -> None:
        if cancel_checker:
            cancel_checker()
        session = _new_session()
        try:
            post_id = str(post.get("post_id") or "")
            stem = post_id or f"post_{idx}"
            user_name = _safe_filename_part(str(post.get("user_name") or "未知作者"))
            post_dir_name = f"{idx:02d}_{user_name}"
            if post_id:
                post_dir_name += f"_{post_id}"
            post_dir = image_dir / post_dir_name
            post_urls = _split_multi_urls(str(post.get("original_image_urls") or ""), sep="|")

            if progress_callback:
                progress_callback(f"下载图片进度 {idx}/{total}: {post_id or '-'}")
            if cancel_checker:
                cancel_checker()

            post_local_paths = _download_url_list(session, post_urls, post_dir, stem=stem, prefix="post")
            post["downloaded_image_count"] = len(post_local_paths)
            post["image_local_paths"] = " | ".join(post_local_paths)

            top_comments = list(post.get("top_comments_data") or [])
            all_comment_local_paths: list[str] = []
            for c_idx, comment in enumerate(top_comments, start=1):
                raw_urls = comment.get("image_urls") or []
                if isinstance(raw_urls, str):
                    comment_urls = _split_multi_urls(raw_urls, sep="|")
                else:
                    comment_urls = _dedup_keep_order(
                        [_to_original_pic_url(str(u)) for u in list(raw_urls) if str(u).strip()]
                    )

                comment_urls = _dedup_image_urls(comment_urls)
                comment_local_paths = _download_url_list(
                    session,
                    comment_urls,
                    post_dir,
                    stem=stem,
                    prefix=f"comment{c_idx}",
                )
                comment["image_urls"] = " | ".join(comment_urls)
                comment["image_local_paths"] = " | ".join(comment_local_paths)
                all_comment_local_paths.extend(comment_local_paths)

            post["top_comments_data"] = top_comments
            post["downloaded_comment_image_count"] = len(all_comment_local_paths)
            post["comment_image_local_paths"] = " | ".join(all_comment_local_paths)
            post["image_local_paths_all"] = " | ".join(post_local_paths + all_comment_local_paths)
            if cancel_checker:
                cancel_checker()
        finally:
            session.close()

    rows = list(posts)
    if not rows:
        return

    worker_count = max(1, min(int(max_workers or 6), len(rows)))
    if worker_count == 1:
        for idx, post in enumerate(rows, start=1):
            if cancel_checker:
                cancel_checker()
            _download_one_post(idx, post, len(rows))
        return

    with ThreadPoolExecutor(max_workers=worker_count, thread_name_prefix="image-download") as executor:
        futures = {
            executor.submit(_download_one_post, idx, post, len(rows)): idx
            for idx, post in enumerate(rows, start=1)
        }
        for future in as_completed(futures):
            if cancel_checker:
                cancel_checker()
            idx = futures[future]
            try:
                future.result()
            except Exception as err:
                if progress_callback:
                    progress_callback(f"下载图片失败 {idx}/{len(rows)}: {type(err).__name__}: {err}")


def build_summary(posts: Iterable[dict]) -> dict:
    return _export_build_summary(posts)


def analyze_active_period(posts: Iterable[dict]) -> dict:
    rows = list(posts)
    hour_counts = [0] * 24
    valid = 0
    for row in rows:
        dt = parse_publish_datetime(str(row.get("publish_time", "") or ""))
        if not dt:
            continue
        hour_counts[dt.hour] += 1
        valid += 1

    if valid == 0:
        return {
            "valid_posts": 0,
            "hour_counts": hour_counts,
            "top_hour": None,
            "top_hour_count": 0,
            "top_two_hour_start": None,
            "top_two_hour_count": 0,
            "low_hour": None,
            "low_hour_count": 0,
            "low_two_hour_start": None,
            "low_two_hour_count": 0,
            "recommended_anchor_hour": 20,
        }

    top_hour = max(range(24), key=lambda h: hour_counts[h])
    low_hour = min(range(24), key=lambda h: hour_counts[h])
    two_hour_scores = [hour_counts[h] + hour_counts[(h + 1) % 24] for h in range(24)]
    top_two_start = max(range(24), key=lambda h: two_hour_scores[h])
    low_two_start = min(range(24), key=lambda h: two_hour_scores[h])

    return {
        "valid_posts": valid,
        "hour_counts": hour_counts,
        "top_hour": top_hour,
        "top_hour_count": hour_counts[top_hour],
        "top_two_hour_start": top_two_start,
        "top_two_hour_count": two_hour_scores[top_two_start],
        "low_hour": low_hour,
        "low_hour_count": hour_counts[low_hour],
        "low_two_hour_start": low_two_start,
        "low_two_hour_count": two_hour_scores[low_two_start],
        "recommended_anchor_hour": top_two_start,
    }


def write_summary_txt(
    summary: dict,
    txt_path: Path,
    leaderboards: dict | None = None,
    active_period: dict | None = None,
    all_posts_summary: dict | None = None,
    carryover_hours: int = 0,
) -> None:
    _export_write_summary_txt(
        summary,
        txt_path,
        leaderboards=leaderboards,
        active_period=active_period,
        all_posts_summary=all_posts_summary,
        carryover_hours=carryover_hours,
        format_leaderboard_line=_format_leaderboard_line,
    )


def _calc_date_distribution_fit(
    all_dist: dict[str, int],
    selected_dist: dict[str, int],
) -> dict:
    return _export_calc_date_distribution_fit(all_dist, selected_dist)


def normalize_date(text: str) -> str | None:
    return _normalize_date(text)


def _date_key_from_publish_time(text: str, parsed_dt: datetime | None = None) -> str:
    if parsed_dt is not None:
        return parsed_dt.strftime("%Y-%m-%d")
    m = re.search(r"(\d{4}-\d{1,2}-\d{1,2})", str(text or "").strip())
    if m:
        return m.group(1)
    return "未知日期"


def _is_nested_feed_item(item: Tag) -> bool:
    parent = item.parent
    while isinstance(parent, Tag):
        if parent.get("action-type") == "feed_list_item":
            return True
        if _is_forwarded_content_container(parent):
            return True
        parent = parent.parent
    return False


def _select_outer(root: Tag, selector: str) -> list[Tag]:
    return [
        node
        for node in root.select(selector)
        if isinstance(node, Tag) and not _is_inside_forwarded_content(node, root)
    ]


def _select_one_outer(root: Tag, selector: str) -> Tag | None:
    for node in _select_outer(root, selector):
        return node
    return None


def _is_inside_forwarded_content(node: Tag, root: Tag | None = None) -> bool:
    current: Tag | None = node if isinstance(node, Tag) else None
    while isinstance(current, Tag):
        if root is not None and current is root:
            return False
        if _is_forwarded_content_container(current):
            return True
        if root is not None and current is not node and current.get("action-type") == "feed_list_item":
            return True
        current = current.parent if isinstance(current.parent, Tag) else None
    return False


def _is_forwarded_content_container(node: Tag) -> bool:
    node_type = str(node.get("node-type") or "").strip()
    if node_type in FORWARDED_NODE_TYPES:
        return True
    classes = node.get("class") or []
    if isinstance(classes, str):
        class_names = classes.split()
    else:
        class_names = [str(item) for item in classes]
    return any(name in FORWARDED_CLASS_NAMES or name.startswith("WB_feed_expand") for name in class_names)


def _extract_user_name(item: Tag) -> str:
    selectors = [
        "a[node-type='feed_list_item_name']",
        "a.W_f14.W_fb.S_txt1[usercard]",
        "a[usercard][nick-name]",
        "a[usercard]",
    ]
    for sel in selectors:
        text = _first_text(_select_outer(item, sel))
        if text:
            return _clean_text(text)
    return ""


def _extract_author_id(item: Tag) -> str:
    tbinfo = str(item.get("tbinfo") or "")
    m = re.search(r"ouid=(\d+)", tbinfo)
    if m:
        return m.group(1)

    anchor = _select_one_outer(item, "a[usercard]")
    if anchor:
        usercard = str(anchor.get("usercard") or "")
        m = re.search(r"id=(\d+)", usercard)
        if m:
            return m.group(1)
    return ""


def _extract_publish_time(item: Tag) -> str:
    date_link = _select_one_outer(item, "a[node-type='feed_list_item_date']")
    if not date_link:
        return ""
    title = date_link.get("title")
    if title:
        return _clean_text(str(title))
    return _clean_text(date_link.get_text(" ", strip=True))


def _extract_content(item: Tag) -> str:
    selectors = [
        "div[node-type='feed_list_content_full']",
        "p[node-type='feed_list_content_full']",
        "div[node-type='feed_list_content']",
        "p[node-type='feed_list_content']",
    ]
    texts: list[str] = []
    for sel in selectors:
        for node in _select_outer(item, sel):
            text = _clean_text_preserve(node.get_text("\n", strip=True))
            if text:
                texts.append(text)
    if not texts:
        return ""
    return _remove_expand_hint(max(texts, key=len))


def _extract_has_video(item: Tag) -> bool:
    selectors = [
        ".WB_video",
        "video",
        "[action-type='feed_list_media_play']",
        "[action-type='feed_list_media_video']",
        "[node-type='fl_h5_video']",
        "a[suda-data*='video']",
    ]
    for selector in selectors:
        if _select_one_outer(item, selector):
            return True
    return False


def _extract_post_url(item: Tag) -> str:
    date_link = _select_one_outer(item, "a[node-type='feed_list_item_date']")
    if not date_link:
        return ""
    href = str(date_link.get("href") or "").strip()
    return _to_absolute_url(href)


def _extract_original_image_urls(item: Tag) -> list[str]:
    urls: list[str] = []

    # 1) 浠庡浘鐗囧尯瀹瑰櫒 action-data 鐨?clear_picSrc 鎻愬彇
    for node in _select_outer(item, ".WB_media_a[action-data], .WB_media_wrap[action-data]"):
        action_data = str(node.get("action-data") or "")
        if "pic" not in action_data.lower():
            continue
        pairs = dict(parse_qsl(action_data))
        clear_pic_src = pairs.get("clear_picSrc") or ""
        urls.extend(
            _to_original_pic_url(candidate)
            for candidate in _split_url_candidates(clear_pic_src)
        )

    # 2) 浠庡浘鐗?media 鑺傜偣鎻愬彇 pid锛屾嫾鍘熷浘 URL
    for media_node in _select_outer(item, "[action-type='feed_list_media_img']"):
        action_data = str(media_node.get("action-data") or "")
        pairs = dict(parse_qsl(action_data))
        pid_text = str(pairs.get("pic_ids") or pairs.get("pid") or "").strip()
        if not pid_text:
            continue
        pids = [x.strip() for x in pid_text.split(",") if x.strip()]
        img = media_node.select_one("img")
        img_src = _to_absolute_url(str((img.get("src") if img else "") or ""))
        ext = _guess_image_ext(img_src)
        host = _extract_sinaimg_host(img_src) or "wx1.sinaimg.cn"
        urls.extend(f"https://{host}/large/{pid}{ext}" for pid in pids)
        if img_src:
            urls.append(_to_original_pic_url(img_src))

    return _dedup_keep_order([u for u in urls if u])


def _to_absolute_url(url: str) -> str:
    return to_absolute_url(url)


def _to_original_pic_url(url: str) -> str:
    return normalize_image_url(url)


def _split_url_candidates(text: str) -> list[str]:
    return _split_multi_urls(text, sep=",")


def _split_multi_urls(text: str, sep: str) -> list[str]:
    raw = str(text or "").strip()
    if not raw:
        return []
    parts = [p.strip() for p in raw.split(sep) if p.strip()]
    return [_to_absolute_url(p) for p in parts]


def _guess_image_ext(url: str) -> str:
    m = re.search(r"(\.(?:jpg|jpeg|png|gif|webp))(?:[?#].*)?$", str(url or ""), flags=re.I)
    if m:
        return m.group(1).lower()
    return ".jpg"


def _extract_sinaimg_host(url: str) -> str:
    m = re.search(r"https?://([^/]*sinaimg\\.cn)/", str(url or ""), flags=re.I)
    return m.group(1) if m else ""


def _dedup_keep_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for v in values:
        if v in seen:
            continue
        seen.add(v)
        out.append(v)
    return out


def _extract_action_count(item: Tag, action_types: list[str]) -> int:
    best = 0
    for action_type in action_types:
        for action in _select_outer(item, f"a[action-type='{action_type}']"):
            text = _clean_text(action.get_text(" ", strip=True))
            best = max(best, parse_count(text))
    return best


def parse_count(text: str) -> int:
    raw = _clean_text(text)
    if not raw:
        return 0
    m = re.search(r"(\d+(?:\.\d+)?)\s*万", raw)
    if m:
        return int(float(m.group(1)) * 10000)
    m = re.search(r"(\d+)", raw)
    if m:
        return int(m.group(1))
    return 0


def _format_comment_for_cell(comment: dict) -> str:
    if not comment:
        return ""
    user = _clean_text(str(comment.get("user_name", "") or "")) or "匿名用户"
    likes = int(comment.get("like_counts", 0) or 0)
    created = _normalize_comment_time(str(comment.get("created_at", "") or ""))
    content = _clean_text(str(comment.get("text", "") or ""))
    has_image = bool(_collect_comment_image_urls(comment))
    if has_image:
        content = _strip_url_like_text(content)
    content = re.sub(r"[\r\n]+", " ", content)
    if not content and has_image:
        content = "（图片评论）"
    if created:
        return f"{user}（赞{likes}，{created}）：{content}"
    return f"{user}（赞{likes}）：{content}"


def _normalize_comment_time(text: str) -> str:
    raw = _clean_text(text)
    if not raw:
        return ""
    try:
        dt = datetime.strptime(raw, "%a %b %d %H:%M:%S %z %Y")
        return dt.strftime("%Y-%m-%d %H:%M")
    except ValueError:
        return raw


def _format_posts_date_range(posts: list[dict]) -> str:
    dates: list[datetime] = []
    for post in posts:
        dt = parse_publish_datetime(str(post.get("publish_time", "") or ""))
        if dt:
            dates.append(dt)
    if not dates:
        today = datetime.now().strftime("%Y-%m-%d")
        return f"{today} 至 {today}"
    start = min(dates).strftime("%Y-%m-%d")
    end = max(dates).strftime("%Y-%m-%d")
    return f"{start} 至 {end}"


def _simplify_hot_comment(text: str) -> str:
    raw = _clean_text(text)
    m = re.match(r"^(.*?)（赞.*?）:\s*(.*)$", raw)
    if m:
        user = _clean_text(m.group(1))
        content = _clean_text(m.group(2))
        return f"{user}：{content}"
    return raw


def _to_rel_path(base_dir: Path, target: Path) -> str:
    try:
        rel = target.resolve().relative_to(base_dir.resolve())
        return str(rel).replace("\\", "/")
    except Exception:
        return str(target.resolve()).replace("\\", "/")


def _report_divider_line() -> str:
    return "─" * 25


def _select_weekly_posts(posts: Iterable[dict], limit: int = 15) -> list[dict]:
    rows = list(posts)
    selected = [row for row in rows if not _should_skip_weekly_post(row)]
    return selected[: max(1, limit)]


def select_weekly_posts(posts: Iterable[dict], limit: int = 15) -> list[dict]:
    return _select_weekly_posts(posts, limit=limit)


def _should_skip_weekly_post(post: dict) -> bool:
    excluded, _reason = should_exclude_post(post)
    return excluded


def _is_video_post(post: dict) -> bool:
    if bool(post.get("has_video")):
        return True
    text = _clean_text(str(post.get("content") or "")).lower()
    post_url = _clean_text(str(post.get("post_url") or "")).lower()
    hit_keyword = any(k in text for k in ("视频", "vid", "播放量"))
    has_video_link = any(k in text or k in post_url for k in ("video.weibo.com", "weibo.com/tv", "/tv/"))
    return hit_keyword and has_video_link


def _is_summary_post(content: str) -> bool:
    raw = _clean_text(content)
    if not raw:
        return False
    low = raw.lower()
    patterns = [
        r"二创精选",
        r"本周精选",
        r"周报",
        r"汇总",
        r"合集",
        r"索引",
        r"导航",
        r"整理了",
        r"发布在.?b站",
        r"前往观赏",
        r"网页链接",
        r"文章在该链接",
    ]
    return any(re.search(p, low, flags=re.I) for p in patterns)


def _get_embed_image_paths(post: dict) -> list[str]:
    all_paths = _split_multi_urls(str(post.get("image_local_paths_all") or ""), sep="|")
    if all_paths:
        return all_paths
    post_paths = _split_multi_urls(str(post.get("image_local_paths") or ""), sep="|")
    comment_paths = _split_multi_urls(str(post.get("comment_image_local_paths") or ""), sep="|")
    return _dedup_keep_order(post_paths + comment_paths)


def _iter_report_comments(post: dict) -> list[dict]:
    top_comments = post.get("top_comments_data") or []
    result: list[dict] = []
    if isinstance(top_comments, list) and top_comments:
        for item in top_comments[:3]:
            if not isinstance(item, dict):
                continue
            image_urls = _split_multi_urls(str(item.get("image_urls") or ""), sep="|")
            image_local_paths = _split_multi_urls(str(item.get("image_local_paths") or ""), sep="|")
            text = _clean_text_preserve(str(item.get("text", "") or ""))
            if image_urls or image_local_paths:
                text = _strip_url_like_text(text)
            result.append(
                {
                    "user_name": _clean_text(str(item.get("user_name", "") or "")),
                    "text": text,
                    "image_urls": " | ".join(image_urls),
                    "image_local_paths": " | ".join(image_local_paths),
                }
            )
    if result:
        return result

    fallback = [
        _clean_text(str(post.get("top_comment_1", "") or "")),
        _clean_text(str(post.get("top_comment_2", "") or "")),
        _clean_text(str(post.get("top_comment_3", "") or "")),
    ]
    return [{"user_name": "", "text": x, "image_urls": "", "image_local_paths": ""} for x in fallback if x]


def _format_hot_comment_text(comment: dict) -> str:
    user = _clean_text(str(comment.get("user_name", "") or ""))
    text = _clean_text_preserve(str(comment.get("text", "") or ""))
    if user and text:
        return f"{user}：{text}"
    if text:
        return text
    if user and _split_multi_urls(str(comment.get("image_local_paths") or ""), sep="|"):
        return f"{user}： （图片评论）"
    return ""


def _collect_top_comment_image_urls(top_comments: list[dict]) -> list[str]:
    urls: list[str] = []
    for comment in top_comments[:3]:
        if not isinstance(comment, dict):
            continue
        urls.extend(_collect_comment_image_urls(comment))
    return _dedup_image_urls(urls)


def _extract_comment_image_urls(comment: dict) -> list[str]:
    return _collect_comment_image_urls(comment)


def _extract_status_image_urls(status: dict) -> list[str]:
    urls: list[str] = []
    if not isinstance(status, dict):
        return urls
    urls.extend(_extract_urls_from_data_node(status.get("pic_infos")))
    urls.extend(_extract_urls_from_data_node(status.get("pic")))
    urls.extend(_extract_urls_from_data_node(status.get("url_struct")))
    urls.extend(_extract_urls_from_data_node(status.get("url_objects")))
    urls.extend(_extract_urls_from_data_node(status.get("mix_media_info")))
    return _dedup_image_urls(urls)


def _collect_comment_image_urls(comment: dict) -> list[str]:
    if not isinstance(comment, dict):
        return []
    urls: list[str] = []
    direct = comment.get("image_urls")
    if isinstance(direct, str):
        urls.extend(_split_multi_urls(direct, sep="|"))
    elif isinstance(direct, list):
        urls.extend([str(x) for x in direct if str(x).strip()])
    for key in ("pic", "pic_infos", "url_struct", "url_objects", "mix_media_info"):
        urls.extend(_extract_urls_from_data_node(comment.get(key)))
    return _dedup_image_urls(urls)


def _dedup_image_urls(urls: list[str]) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for raw in urls:
        url = _to_original_pic_url(str(raw))
        if not _looks_like_image_url(url):
            continue
        sig = _image_signature(url)
        if sig in seen:
            continue
        seen.add(sig)
        out.append(url)
    return out


def _image_signature(url: str) -> str:
    clean = _to_absolute_url(url).split("?", 1)[0].split("#", 1)[0]
    m = re.search(r"/([A-Za-z0-9]+)\.(?:jpg|jpeg|png|gif|webp)$", clean, flags=re.I)
    if m:
        return m.group(1).lower()
    return clean.lower()


def _extract_urls_from_data_node(node) -> list[str]:
    urls: list[str] = []
    if isinstance(node, str):
        maybe = _to_absolute_url(node)
        if _looks_like_image_url(maybe):
            urls.append(maybe)
        return urls
    if isinstance(node, list):
        for item in node:
            urls.extend(_extract_urls_from_data_node(item))
        return urls
    if isinstance(node, dict):
        for key in ("url", "ori_url", "pic", "pic_url", "thumbnail_pic", "bmiddle_pic", "original_pic"):
            value = node.get(key)
            if isinstance(value, str):
                maybe = _to_absolute_url(value)
                if _looks_like_image_url(maybe):
                    urls.append(maybe)
            elif isinstance(value, dict):
                nested = value.get("url")
                if isinstance(nested, str):
                    maybe = _to_absolute_url(nested)
                    if _looks_like_image_url(maybe):
                        urls.append(maybe)
        for key in ("large", "largest", "orj360", "mw2000", "mw690"):
            value = node.get(key)
            if isinstance(value, dict):
                maybe = _to_absolute_url(str(value.get("url") or ""))
                if _looks_like_image_url(maybe):
                    urls.append(maybe)
        for value in node.values():
            if isinstance(value, (dict, list)):
                urls.extend(_extract_urls_from_data_node(value))
    return urls


def _looks_like_image_url(url: str) -> bool:
    u = _to_absolute_url(url)
    if not u:
        return False
    if re.search(r"\.(?:jpg|jpeg|png|gif|webp)(?:[?#].*)?$", u, flags=re.I):
        return True
    return "sinaimg.cn" in u.lower()


def _html_to_text(text: str) -> str:
    return strip_html_text(text)


def _html_to_text_preserve(text: str) -> str:
    return strip_html_text(text, preserve_newlines=True)


def _strip_url_like_text(text: str) -> str:
    raw = _clean_text(text)
    raw = re.sub(r"https?://\S+", " ", raw, flags=re.I)
    raw = re.sub(r"\b(?:t\.cn|weibo\.cn|weibo\.com)/\S+", " ", raw, flags=re.I)
    raw = re.sub(r"(网页链接|网页链接\:?)", " ", raw, flags=re.I)
    raw = re.sub(r"\s{2,}", " ", raw)
    return raw.strip()


def _strip_specific_topic_tags(text: str) -> str:
    return clean_topic_tags(text)


def _strip_specific_topic_tags_preserve(text: str) -> str:
    return clean_topic_tags(text, preserve_newlines=True)


def _safe_filename_part(text: str, max_len: int = 24) -> str:
    raw = _clean_text(text) or "item"
    raw = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", raw)
    raw = re.sub(r"\s+", "_", raw).strip("._ ")
    return (raw or "item")[:max_len]


def _clean_report_text(text: str) -> str:
    raw = _strip_specific_topic_tags_preserve(_clean_text_preserve(text))
    raw = _replace_weibo_emoticons(raw)
    raw = _replace_unicode_emoji(raw)
    # 去掉微博字体图标的私有区字符、零宽字符，避免周报里出现乱码方块。
    raw = re.sub(r"[\u200b-\u200f\u202a-\u202e\ufeff]", "", raw)
    raw = raw.replace("\ufe0f", "")
    raw = re.sub(r"[\ue000-\uf8ff]", "", raw)
    return _clean_text_preserve(raw)


def _replace_weibo_emoticons(text: str) -> str:
    mapping = {
        # 开心/友好
        "抱一抱": "(づ｡◕‿‿◕｡)づ",
        "抱抱": "(づ￣ 3￣)づ",
        "打call": "ヾ(≧▽≦*)o",
        "哈哈": "(๑>◡<๑)",
        "嘻嘻": "(*^▽^*)",
        "可爱": "(=^･ω･^=)",
        "爱你": "( ˘ ³˘)♥",
        "亲亲": "( ˘ ³˘)♥",
        "鼓掌": "(*'ω'ﾉﾉﾞ☆",
        "送花花": "(✿◡‿◡)",
        "赞": "(๑•̀ㅂ•́)و",
        "ok": "(๑•̀ㅂ•́)و",

        # 笑哭/偏搞笑（和“哭泣”区分）
        "笑cry": "(≧▽≦;)",
        "笑哭": "(≧▽≦;)",
        "偷笑": "(￣▽￣)~*",
        "憨笑": "(≧∀≦)ゞ",
        "doge": "(￣▽￣)",
        "doge脸": "(￣▽￣)",
        "二哈": "(哈▽哈)",
        "允悲": "(；▽；)",

        # 难过/哭泣
        "泪": "(T_T)",
        "流泪": "(；﹏；)",
        "泪奔": "(ಥ_ಥ)",
        "悲伤": "(Q_Q)",
        "大哭": "(╥﹏╥)",
        "委屈": "(｡•́︿•̀｡)",
        "可怜": "(´；ω；`)",

        # 其他常见情绪
        "思考": "( •̀ .̫ •́ )",
        "疑问": "( ?_? )",
        "跪了": "_(:3」∠)_",
        "馋嘴": "(๑´ڡ`๑)",
        "干饭人": "(๑´ڡ`๑)",
        "裂开": "(⊙_⊙;)",
        "苦涩": "(＞﹏＜)",
        "哇": "(✧ω✧)",
        "心": "<3",
        "给你小心心": "<3<3",
    }

    def repl(match: re.Match[str]) -> str:
        key = _clean_text(match.group(1))
        return mapping.get(key, "(๑•ᴗ•๑)")

    return re.sub(r"\[([^\[\]]{1,24})\]", repl, text)


def _replace_unicode_emoji(text: str) -> str:
    emoji_map = {
        # 开心/友好
        "😀": "(*^_^*)",
        "😄": "(*^o^*)",
        "😁": "(๑>◡<๑)",
        "😆": "(≧ω≦)",
        "😊": "(*^_^*)",
        "😍": "(=^.^=)",
        "🥰": "(=^.^=)",
        "😘": "(*^3^*)",
        "😋": "(^q^)",
        "🤗": "(*^_^*)",
        "😇": "(^_^)v",
        "😎": "B-)",

        # 笑哭（和哭泣区分）
        "😂": "(≧▽≦;)",
        "🤣": "xD",
        "😹": "(=^▽^=;)",
        "😅": "(*^_^*;)",

        # 难过/哭泣
        "😢": "(；﹏；)",
        "😭": "(╥﹏╥)",
        "🥲": "(´；ω；`)",
        "😿": "(T_T)",
        "🥺": "(｡•́︿•̀｡)",

        # 其他
        "🤔": "( •̀ .̫ •́ )",
        "😴": "(-_-) zzz",
        "😐": "( -_- )",
        "😑": "( -_- )",
        "😶": "( ._. )",
        "❤": "<3",
        "❤️": "<3",
        "💗": "<3",
        "💖": "<3",
        "💘": "<3",
        "💕": "<3",
        "💞": "<3",
        "✨": "(*_*)",
        "🌟": "(*_*)",
        "👍": "(^_^)v",
        "👏": "(*'ω'ﾉﾉﾞ☆",
        "🙏": "(*^_^*)",
    }
    out = text
    for emo, face in emoji_map.items():
        out = out.replace(emo, face)
    # 兜底替换仍残留的 emoji 为温和颜文字，避免跨平台显示差异
    return re.sub(r"[\U0001F300-\U0001FAFF\u2600-\u27BF]", "(๑•ᴗ•๑)", out)


def _set_east_asia_font(style, font_name: str) -> None:
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _set_run_east_asia_font(run, font_name: str) -> None:
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _first_text(elements: list) -> str:
    if not elements:
        return ""
    return elements[0].get_text(" ", strip=True)


def _clean_text(text: str) -> str:
    return normalize_weibo_text(text)


def _clean_text_preserve(text: str) -> str:
    return collapse_blank_lines(text)

