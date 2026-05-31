from __future__ import annotations

from collections.abc import Iterable
from contextlib import suppress
from pathlib import Path
from typing import Any

from docx import Document as create_doc
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from export.context import ExportContext
from export.docx_images import add_comment_images, add_post_images
from export.docx_splitter import DOCX_SIZE_LIMIT_BYTES, numbered_docx_path
from export.docx_styles import (
    add_heading,
    add_hyperlink,
    add_paragraph_text,
    add_run_text,
    setup_document_styles,
)
from export.report_helpers import (
    clean_report_text,
    format_hot_comment_text,
    format_posts_date_range,
    iter_report_comments,
    select_weekly_posts,
)


DEFAULT_REPORT_TITLE = "warma超话周报"


def export_docx(ctx: ExportContext, path: Path | None = None, max_bytes: int = DOCX_SIZE_LIMIT_BYTES) -> list[Path]:
    return export_weekly_report_docx(
        ctx.selected_posts,
        path or ctx.run_dir / "weekly_report.docx",
        title=str(ctx.config.get("report_title") or DEFAULT_REPORT_TITLE),
        leaderboards=(ctx.config.get("leaderboards") if isinstance(ctx.config, dict) else None),
        preselected=True,
        max_bytes=max_bytes,
        ctx=ctx,
    )


def export_weekly_report_docx(
    posts: Iterable[dict[str, Any]],
    docx_path: Path,
    title: str = DEFAULT_REPORT_TITLE,
    leaderboards: dict[str, Any] | None = None,
    preselected: bool = False,
    max_bytes: int = DOCX_SIZE_LIMIT_BYTES,
    ctx: ExportContext | None = None,
) -> list[Path]:
    all_posts = list(posts)
    rows = all_posts[:15] if preselected else select_weekly_posts(all_posts, limit=15)
    export_ctx = ctx or ExportContext(docx_path.parent, rows, all_posts, {"report_title": title}, {})
    board = leaderboards or _build_comment_leaderboards(all_posts)
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    if not rows:
        out_path = numbered_docx_path(docx_path, 1)
        document = build_docx_document(export_ctx, [], title=title, leaderboards=board, include_leaderboards=True)
        document.save(str(out_path))
        return [out_path]

    parts: list[Path] = []
    current: list[dict[str, Any]] = []
    trial_path = docx_path.with_name(f"_{docx_path.stem}_trial{docx_path.suffix}")
    limit = max(1, int(max_bytes or DOCX_SIZE_LIMIT_BYTES))

    for post in rows:
        trial_rows = current + [post]
        document = build_docx_document(
            export_ctx,
            trial_rows,
            title=title,
            leaderboards=board,
            include_leaderboards=(len(parts) == 0),
        )
        document.save(str(trial_path))
        trial_size = trial_path.stat().st_size if trial_path.exists() else 0
        with suppress(Exception):
            trial_path.unlink(missing_ok=True)

        if trial_size > limit and current:
            out_path = numbered_docx_path(docx_path, len(parts) + 1)
            document = build_docx_document(
                export_ctx,
                current,
                title=title,
                leaderboards=board,
                include_leaderboards=(len(parts) == 0),
            )
            document.save(str(out_path))
            parts.append(out_path)
            current = [post]
        else:
            current = trial_rows

    if current:
        out_path = numbered_docx_path(docx_path, len(parts) + 1)
        document = build_docx_document(
            export_ctx,
            current,
            title=title,
            leaderboards=board,
            include_leaderboards=(len(parts) == 0),
        )
        document.save(str(out_path))
        parts.append(out_path)
    return parts


def export_weekly_report_sum_docx(
    posts: Iterable[dict[str, Any]],
    docx_path: Path,
    title: str = DEFAULT_REPORT_TITLE,
    leaderboards: dict[str, Any] | None = None,
    preselected: bool = False,
    ctx: ExportContext | None = None,
) -> Path:
    all_posts = list(posts)
    rows = all_posts[:15] if preselected else select_weekly_posts(all_posts, limit=15)
    export_ctx = ctx or ExportContext(docx_path.parent, rows, all_posts, {"report_title": title}, {})
    board = leaderboards or _build_comment_leaderboards(all_posts)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document = build_docx_document(export_ctx, rows, title=title, leaderboards=board, include_leaderboards=True)
    document.save(str(docx_path))
    return docx_path


def build_docx_document(
    ctx: ExportContext,
    posts_chunk: list[dict[str, Any]] | None = None,
    title: str | None = None,
    leaderboards: dict[str, Any] | None = None,
    include_leaderboards: bool = True,
):
    document = create_doc()
    setup_document_styles(document)
    rows = list(posts_chunk if posts_chunk is not None else ctx.selected_posts)
    if include_leaderboards:
        _render_docx_header(
            document,
            rows,
            title or str(ctx.config.get("report_title") or DEFAULT_REPORT_TITLE),
            leaderboards or {},
        )
    for rank, post in enumerate(rows, start=1):
        render_docx_post(document, post, rank, ctx)
    return document


def render_docx_post(document, post: dict[str, Any], rank: int, ctx: ExportContext) -> None:
    author = _clean_inline_text(str(post.get("user_name") or "未知作者"))
    content = clean_report_text(str(post.get("content") or "")) or "（无正文）"
    publish_time = _clean_inline_text(str(post.get("publish_time") or ""))
    post_url = _clean_inline_text(str(post.get("post_url") or ""))

    body = document.add_paragraph()
    body.paragraph_format.space_after = Pt(6)
    add_run_text(body, f"@{author}：", size=12, bold=True)
    add_run_text(body, content, size=12, bold=False)

    if publish_time:
        add_paragraph_text(document, f"发送时间：{publish_time}", size=10, color=(73, 80, 87))

    add_post_images(document, post, ctx)

    comments = iter_report_comments(post)
    if comments:
        add_paragraph_text(document, "热评：", size=10, bold=True, color=(90, 98, 104))
        for comment in comments:
            comment_text = clean_report_text(format_hot_comment_text(comment))
            if comment_text:
                add_paragraph_text(document, comment_text, size=9, color=(108, 117, 125))
            add_comment_images(document, comment, ctx)

    if post_url:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        add_run_text(paragraph, "原帖链接：", size=10, color=(73, 80, 87))
        add_hyperlink(paragraph, post_url, post_url)

    document.add_paragraph("")


def _render_docx_header(document, rows: list[dict[str, Any]], title: str, board: dict[str, Any]) -> None:
    add_heading(document, title, level=1)
    add_paragraph_text(
        document,
        f"帖子选取日期：{format_posts_date_range(rows)}",
        size=11,
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
    )

    add_heading(document, "本周社区互动榜", level=2)
    add_paragraph_text(document, "评论数量榜 Top3", size=11, bold=True)
    count_rows = list(board.get("comment_count_top3") or [])
    if count_rows:
        for item in count_rows:
            add_paragraph_text(
                document,
                _format_leaderboard_line(item, include_hot=False, include_like_total=False, include_post_span=True),
                size=10,
            )
    else:
        add_paragraph_text(document, "暂无评论数据", size=10)

    add_paragraph_text(document, "评论质量榜 Top3", size=11, bold=True)
    quality_rows = list(board.get("comment_quality_top3") or [])
    if quality_rows:
        for item in quality_rows:
            add_paragraph_text(document, _format_leaderboard_line(item, include_hot=True), size=10)
    else:
        add_paragraph_text(document, "暂无评论数据", size=10)

    add_heading(document, "本周热帖Top15", level=2)


def _build_comment_leaderboards(posts: list[dict[str, Any]]) -> dict[str, Any]:
    from modules.comments.ranking import build_comment_leaderboards

    return build_comment_leaderboards(posts, top_n=3)


def _format_leaderboard_line(
    item: dict[str, Any],
    include_hot: bool = True,
    include_like_total: bool = True,
    include_post_span: bool = False,
) -> str:
    user_name = _clean_inline_text(str(item.get("user_name") or "匿名用户"))
    line = f"{int(item.get('rank') or 0)}. @{user_name}：评论 {int(item.get('comment_count') or 0)} 条"
    if include_post_span:
        line += f"，评论过 {int(item.get('commented_post_count') or 0)} 条帖子"
    elif include_like_total:
        line += f"，本周评论获赞 {int(item.get('comment_likes_total') or 0)}"
    if include_hot:
        line += f"，热评前三 {int(item.get('hot_top3_count') or 0)} 次"
    return line


def _clean_inline_text(text: str) -> str:
    return " ".join(str(text or "").replace("\r\n", "\n").replace("\r", "\n").split())
