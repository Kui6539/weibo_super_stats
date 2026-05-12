from __future__ import annotations

from pathlib import Path
from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_post_images(document: DocxDocument, post: dict[str, Any], ctx: Any) -> None:
    for image_path in _split_paths(post.get("image_local_paths")):
        _add_scaled_right_aligned_picture(document, image_path, get_image_display_width("post_image", ctx), ctx)


def add_comment_images(document: DocxDocument, comment: dict[str, Any], ctx: Any) -> None:
    for image_path in _split_paths(comment.get("image_local_paths")):
        _add_scaled_right_aligned_picture(document, image_path, get_image_display_width("comment_image", ctx), ctx)


def get_image_display_width(image_type: str, ctx: Any) -> float:
    return 0.25 if image_type == "comment_image" else 0.50


def _add_scaled_right_aligned_picture(document: DocxDocument, image_path: str, scale: float, ctx: Any) -> None:
    path = Path(str(image_path))
    if not path.exists():
        if hasattr(ctx, "warnings"):
            ctx.warnings.append(f"DOCX 图片缺失：{image_path}")
        return
    try:
        section = document.sections[-1]
        page_width = int(section.page_width or 0)
        left_margin = int(section.left_margin or 0)
        right_margin = int(section.right_margin or 0)
        usable_width = page_width - left_margin - right_margin
        width = int(usable_width * max(0.05, min(1.0, float(scale))))
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        paragraph.add_run().add_picture(str(path), width=width)
    except Exception as err:
        if hasattr(ctx, "warnings"):
            ctx.warnings.append(f"DOCX 图片插入失败：{type(err).__name__}")


def _split_paths(value: Any) -> list[str]:
    if isinstance(value, list):
        parts = [str(item).strip() for item in value]
    else:
        parts = [part.strip() for part in str(value or "").replace("\n", "|").split("|")]
    return [part for part in parts if part]
