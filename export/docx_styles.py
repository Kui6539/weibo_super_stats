from __future__ import annotations

from typing import Any

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_NAME = "Microsoft YaHei"


def setup_document_styles(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    _set_east_asia_font(normal, FONT_NAME)

    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_heading(document, text: str, level: int = 1):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if level == 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    run = add_run_text(
        paragraph,
        str(text),
        size=26 if level == 1 else 14 if level == 2 else 11,
        bold=True,
        color=(33, 37, 41),
    )
    return paragraph


def add_paragraph_text(
    document,
    text: str,
    size: int = 11,
    bold: bool = False,
    color: tuple[int, int, int] | None = None,
    alignment: WD_PARAGRAPH_ALIGNMENT | None = None,
):
    paragraph = document.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(4)
    add_run_text(paragraph, str(text or ""), size=size, bold=bold, color=color)
    return paragraph


def add_run_text(
    paragraph,
    text: str,
    size: int = 11,
    bold: bool = False,
    color: tuple[int, int, int] | None = None,
):
    run = paragraph.add_run()
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    if color:
        run.font.color.rgb = RGBColor(*color)
    _set_run_east_asia_font(run, FONT_NAME)
    _add_preserved_text(run, str(text or ""))
    return run


def add_hyperlink(paragraph, text: str, url: str):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = paragraph.add_run(str(text))
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.underline = True
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    _set_run_east_asia_font(run, FONT_NAME)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)
    return run


def _add_preserved_text(run, text: str) -> None:
    parts = str(text or "").splitlines()
    if not parts:
        run.add_text("")
        return
    for idx, part in enumerate(parts):
        if idx > 0:
            run.add_break()
        if part:
            run.add_text(part)


def _set_east_asia_font(style: Any, font_name: str) -> None:
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_run_east_asia_font(run: Any, font_name: str) -> None:
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
